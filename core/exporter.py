"""
Exporter Module - Export translated chapters to .docx
Converts markdown translation output to a formatted Word document.
Supports single-chapter and batch (merged) export with optional .docx template.
"""

import io
import re
import sys
from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
from .extractor import strip_glossary_table
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_FILENAME = 'TemplateNew.docx'

# Unicode ranges not reliably covered by Times New Roman (or any single Latin font).
# These need an explicit fallback so Word doesn't corrupt them.
_UNICODE_FALLBACK_RANGES = (
    (0x2030, 0x206F),  # General Punctuation extras: ※ ‼ ⁇ etc. (em/en dash & quotes are below 0x2030, safe)
    (0x2190, 0x23FF),  # Arrows + Supplemental Arrows + Misc Technical
    (0x2400, 0x27BF),  # Box Drawing / Geometric Shapes / Misc Symbols / Dingbats
    (0x2E80, 0x2EFF),  # CJK Radicals Supplement
    (0x3000, 0x303F),  # CJK Symbols and Punctuation
    (0x3040, 0x30FF),  # Hiragana / Katakana
    (0x4E00, 0x9FFF),  # CJK Unified Ideographs
    (0xAC00, 0xD7AF),  # Hangul Syllables
    (0xFF00, 0xFFEF),  # Halfwidth / Fullwidth Forms
)
# A single Unicode-capable font available on Windows 8+ and all Office installs.
_FALLBACK_FONT = 'Segoe UI Symbol'


def _char_needs_fallback(ch: str) -> bool:
    cp = ord(ch)
    return any(lo <= cp <= hi for lo, hi in _UNICODE_FALLBACK_RANGES)


def _split_for_font(text: str):
    """Yield (segment, needs_fallback_font) pairs."""
    if not text:
        return
    cur = [text[0]]
    cur_fallback = _char_needs_fallback(text[0])
    for ch in text[1:]:
        fb = _char_needs_fallback(ch)
        if fb == cur_fallback:
            cur.append(ch)
        else:
            yield ''.join(cur), cur_fallback
            cur = [ch]
            cur_fallback = fb
    yield ''.join(cur), cur_fallback


def _add_run(paragraph, text: str, bold: bool = False, italic: bool = False,
             base_font: Optional[str] = 'Times New Roman'):
    """Add text as one or more runs, switching to fallback font for Unicode symbols.
    Pass base_font=None to inherit font from the paragraph style (e.g. when using a template)."""
    for segment, needs_fallback in _split_for_font(text):
        run = paragraph.add_run(segment)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if needs_fallback:
            run.font.name = _FALLBACK_FONT
        elif base_font is not None:
            run.font.name = base_font


def _find_template() -> Optional[Path]:
    """Locate TemplateNew.docx from several candidate directories."""
    from .storage.config import get_app_data_dir
    candidates = [
        get_app_data_dir() / TEMPLATE_FILENAME,
        Path(__file__).parent.parent / TEMPLATE_FILENAME,   # project root (dev)
        Path.cwd() / TEMPLATE_FILENAME,
    ]
    if getattr(sys, 'frozen', False):
        candidates.insert(1, Path(sys.executable).parent / TEMPLATE_FILENAME)
    for p in candidates:
        if p.exists():
            return p
    return None


def _open_template_doc() -> tuple:
    """
    Returns (doc, used_template).
    Opens TemplateNew.docx and clears its body content while keeping styles/page setup.
    Falls back to a fresh Document if template is not found.
    """
    tpl_path = _find_template()
    if tpl_path:
        doc = Document(str(tpl_path))
        # Remove all body children except w:sectPr (section/page properties)
        body = doc.element.body
        for child in list(body):
            if child.tag != qn('w:sectPr'):
                body.remove(child)
        return doc, True
    return Document(), False


def _set_doc_margins(doc: Document):
    """Set comfortable reading margins."""
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)


def _configure_normal_style(doc: Document):
    """Configure Normal style: serif font, comfortable line spacing."""
    style = doc.styles['Normal']
    font  = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(20)


def _add_inline_content(paragraph, text: str, base_font: Optional[str] = 'Times New Roman'):
    """
    Parse inline markdown and add runs to paragraph.
    Handles: ***bold-italic***, **bold**, *italic*, plain text.
    Each run is split by character type so Unicode symbols get a proper fallback font.
    Pass base_font=None to inherit font from paragraph style (template use).
    """
    pattern = r'(\*{3}.+?\*{3}|\*{2}.+?\*{2}|\*.+?\*)'
    parts = re.split(pattern, text, flags=re.DOTALL)

    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***') and len(part) > 6:
            _add_run(paragraph, part[3:-3], bold=True, italic=True, base_font=base_font)
        elif part.startswith('**') and part.endswith('**') and len(part) > 4:
            _add_run(paragraph, part[2:-2], bold=True, base_font=base_font)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            _add_run(paragraph, part[1:-1], italic=True, base_font=base_font)
        else:
            _add_run(paragraph, part, base_font=base_font)


def _is_scene_break(text: str) -> bool:
    """Return True if the line is a scene-break marker (*** / --- / ___)."""
    return bool(re.match(r'^(\*{3}|-{3,}|_{3,})$', text.strip()))


def _add_scene_break(doc: Document):
    """Add a centered *** paragraph as a scene break."""
    p = doc.add_paragraph('* * *')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_page_break(doc: Document):
    """Insert a hard page break as an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def _add_chapter_to_doc(
    doc: Document,
    series: Dict[str, Any],
    chapter: Dict[str, Any],
    show_series_subtitle: bool = True,
    use_template_styles: bool = False,
):
    """
    Append a single chapter's content to an existing Document.
    If use_template_styles=True, only applies font overrides when the template
    doesn't already define them (avoids clobbering template Heading styles).
    """
    series_title  = series.get('title', '')
    chapter_num   = chapter.get('chapterNumber', 1)
    chapter_title = chapter.get('title', '').strip()
    content       = strip_glossary_table(chapter.get('translatedContent', ''))

    heading_text = f'Bab {chapter_num}'
    if chapter_title:
        heading_text += f' — {chapter_title}'

    # Chapter heading — use Heading 1 style (from template or default)
    h = doc.add_heading(heading_text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not use_template_styles:
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)

    # Series subtitle (italic, grey)
    if series_title and show_series_subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(18)
        subtitle_font = None if use_template_styles else 'Times New Roman'
        _add_run(sub, series_title, italic=True, base_font=subtitle_font)
        if not use_template_styles:
            for run in sub.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # blank spacer after header

    # When using template, inherit body font from template's Normal style.
    # When not using template, explicitly set Times New Roman.
    body_base_font = None if use_template_styles else 'Times New Roman'

    # Body content — split on blank lines
    blocks = re.split(r'\n{2,}', content.strip())
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        heading_match = re.match(r'^(#{1,4})\s+(.+)$', block)
        if heading_match:
            level = min(len(heading_match.group(1)) + 1, 4)
            text  = heading_match.group(2).strip()
            h2 = doc.add_heading(text, level=level)
            if not use_template_styles:
                for run in h2.runs:
                    run.font.name = 'Times New Roman'
            continue

        if _is_scene_break(block):
            _add_scene_break(doc)
            continue

        lines = block.split('\n')
        if len(lines) == 1:
            p = doc.add_paragraph()
            _add_inline_content(p, block, base_font=body_base_font)
        else:
            p = doc.add_paragraph()
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                if i > 0:
                    p.add_run('\n')
                _add_inline_content(p, line, base_font=body_base_font)


def export_chapter_to_docx(series: Dict[str, Any], chapter: Dict[str, Any]) -> bytes:
    """Export a single chapter to .docx bytes (no template)."""
    doc = Document()
    _set_doc_margins(doc)
    _configure_normal_style(doc)
    _add_chapter_to_doc(doc, series, chapter, use_template_styles=False)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_series_to_docx(series: Dict[str, Any], chapters: List[Dict[str, Any]]) -> tuple:
    """
    Export all provided chapters into a single merged .docx.
    Uses TemplateNew.docx if found; falls back to default styling.
    Each chapter starts on a new page.

    Returns:
        (bytes, template_used: bool)
    """
    doc, used_template = _open_template_doc()

    if not used_template:
        _set_doc_margins(doc)
        _configure_normal_style(doc)

    for i, chapter in enumerate(chapters):
        if i > 0:
            _add_page_break(doc)
        _add_chapter_to_doc(doc, series, chapter, use_template_styles=used_template)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), used_template
